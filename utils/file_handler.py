"""
RadiAI.Care 文件處理模塊
統一處理各種文件格式的文本提取
"""

import fitz  # PyMuPDF for PDF
import docx  # python-docx for Word documents
import io
import logging
from typing import Optional, Tuple, Dict, Any
from config.settings import AppConfig

logger = logging.getLogger(__name__)

class FileHandler:
    """文件處理類，支持多種格式的文本提取"""
    
    def __init__(self):
        self.config = AppConfig()
        self.supported_types = self.config.SUPPORTED_FILE_TYPES
        self.max_size_mb = self.config.FILE_SIZE_LIMIT_MB
    
    def validate_file(self, uploaded_file) -> Tuple[bool, str, Dict[str, Any]]:
        """
        驗證上傳文件的有效性
        
        Args:
            uploaded_file: Streamlit上傳的文件對象
            
        Returns:
            Tuple[bool, str, Dict]: (是否有效, 錯誤信息, 文件信息)
        """
        if uploaded_file is None:
            return False, "沒有選擇文件", {}
        
        # 獲取文件信息
        file_info = self._get_file_info(uploaded_file)
        
        # 檢查文件大小
        if file_info['size_mb'] > self.max_size_mb:
            return False, f"文件過大，請上傳小於{self.max_size_mb}MB的文件", file_info
        
        # 檢查文件類型
        if file_info['extension'] not in self.supported_types:
            return False, f"不支持的文件格式，請上傳{', '.join(self.supported_types).upper()}文件", file_info
        
        return True, "", file_info
    
    def extract_text(self, uploaded_file) -> Tuple[Optional[str], Dict[str, Any]]:
        """
        從上傳的文件中提取文本
        
        Args:
            uploaded_file: Streamlit上傳的文件對象
            
        Returns:
            Tuple[Optional[str], Dict]: (提取的文本內容, 處理信息)
        """
        # 先驗證文件
        is_valid, error_msg, file_info = self.validate_file(uploaded_file)
        
        if not is_valid:
            logger.error(f"File validation failed: {error_msg}")
            return None, {'error': error_msg, 'file_info': file_info}
        
        file_extension = file_info['extension']
        
        try:
            # 根據文件類型調用相應的提取方法
            if file_extension == 'txt':
                text = self._extract_from_txt(uploaded_file)
            elif file_extension == 'pdf':
                text = self._extract_from_pdf(uploaded_file)
            elif file_extension in ['docx', 'doc']:
                text = self._extract_from_docx(uploaded_file)
            else:
                logger.error(f"Unsupported file type: {file_extension}")
                return None, {'error': f"不支持的文件類型: {file_extension}", 'file_info': file_info}
            
            # 處理提取結果
            if text and text.strip():
                result_info = {
                    'success': True,
                    'file_info': file_info,
                    'extracted_length': len(text),
                    'preview': text[:200] + "..." if len(text) > 200 else text
                }
                return text.strip(), result_info
            else:
                return None, {
                    'error': '文件中沒有提取到有效文本內容',
                    'file_info': file_info
                }
                
        except Exception as e:
            logger.error(f"Text extraction failed for {file_extension}: {e}")
            return None, {
                'error': f"文本提取失敗: {str(e)}",
                'file_info': file_info
            }
    
    def _get_file_info(self, uploaded_file) -> Dict[str, Any]:
        """獲取文件基本信息"""
        try:
            file_size_bytes = len(uploaded_file.getvalue())
            file_size_mb = file_size_bytes / (1024 * 1024)
            file_extension = uploaded_file.name.lower().split('.')[-1] if '.' in uploaded_file.name else ''
            
            return {
                'name': uploaded_file.name,
                'extension': file_extension,
                'size_bytes': file_size_bytes,
                'size_mb': round(file_size_mb, 2),
                'size_kb': round(file_size_bytes / 1024, 1)
            }
        except Exception as e:
            logger.error(f"Error getting file info: {e}")
            return {
                'name': getattr(uploaded_file, 'name', 'unknown'),
                'extension': 'unknown',
                'size_bytes': 0,
                'size_mb': 0,
                'size_kb': 0
            }
    
    def _extract_from_txt(self, uploaded_file) -> str:
        """從TXT文件提取文本"""
        encodings = ['utf-8', 'gbk', 'big5', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                uploaded_file.seek(0)
                content = uploaded_file.read().decode(encoding)
                logger.info(f"Successfully decoded TXT file with {encoding} encoding")
                return content.strip()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.warning(f"Error with {encoding} encoding: {e}")
                continue
        
        raise ValueError("無法用任何支持的編碼解碼文本文件")
    
    def _extract_from_pdf(self, uploaded_file) -> str:
        """從PDF文件提取文本"""
        try:
            pdf_document = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            text_parts = []
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                page_text = page.get_text()
                
                if page_text.strip():
                    text_parts.append(page_text.strip())
                    
            pdf_document.close()
            
            if text_parts:
                full_text = "\n\n".join(text_parts)
                logger.info(f"Successfully extracted text from {len(text_parts)} PDF pages")
                return full_text
            else:
                raise ValueError("PDF文件中沒有可提取的文本內容")
                
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise ValueError(f"PDF文件處理失敗: {str(e)}")
    
    def _extract_from_docx(self, uploaded_file) -> str:
        """從DOCX文件提取文本"""
        try:
            doc = docx.Document(io.BytesIO(uploaded_file.read()))
            text_parts = []
            
            # 提取段落文本
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
                    paragraph_count += 1
            
            # 提取表格內容
            table_count = 0
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_cells.append(cell.text.strip())
                    if row_cells:
                        table_rows.append(" | ".join(row_cells))
                
                if table_rows:
                    text_parts.extend(table_rows)
                    table_count += 1
            
            if text_parts:
                full_text = "\n".join(text_parts)
                logger.info(f"Successfully extracted text from DOCX: {paragraph_count} paragraphs, {table_count} tables")
                return full_text
            else:
                raise ValueError("DOCX文件中沒有可提取的文本內容")
                
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise ValueError(f"DOCX文件處理失敗: {str(e)}")
    
    def get_supported_formats_info(self) -> Dict[str, str]:
        """獲取支持格式的說明信息"""
        return {
            'pdf': "📄 PDF - 掃描或電子版報告",
            'txt': "📝 TXT - 純文字報告",
            'docx': "📑 DOCX - Word文檔報告"
        }
    
    def estimate_processing_time(self, file_size_mb: float) -> str:
        """估算文件處理時間"""
        if file_size_mb < 1:
            return "< 5秒"
        elif file_size_mb < 5:
            return "5-15秒"
        else:
            return "15-30秒"
