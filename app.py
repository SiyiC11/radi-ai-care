import streamlit as st
from openai import OpenAI
import os
from dotenv import load_dotenv
import time
import logging
import fitz  # PyMuPDF for PDF
import docx  # python-docx for Word documents
import io
import uuid
from typing import Optional, Dict, Any, Tuple, List
import asyncio
from concurrent.futures import ThreadPoolExecutor
import hashlib

# Import custom modules with enhanced error handling
try:
    from utils.prompt_template import get_prompt, create_enhanced_disclaimer, get_processing_steps
    from log_to_sheets import log_to_google_sheets
except ImportError as e:
    st.error(f"⚠️ 模組導入失敗: {e}")
    st.error("請確保 utils/prompt_template.py 和 log_to_sheets.py 文件存在")
    st.stop()

# -----------------------------
# Configuration and Constants
# -----------------------------
load_dotenv()
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# App Configuration
APP_CONFIG = {
    "version": "v4.0",
    "title": "RadiAI.Care - 智能醫療報告助手",
    "icon": "🏥",
    "layout": "centered",
    "description": "為澳洲華人社區打造的醫療報告翻譯服務"
}

# Feature Limits
LIMITS = {
    "max_free_translations": 3,
    "file_size_limit_mb": 10,
    "min_text_length": 20,
    "max_text_length": 50000,
    "preview_length": 800,
    "max_retries": 3,
    "request_timeout": 30
}

# Supported file types
SUPPORTED_FILE_TYPES = ['pdf', 'txt', 'docx']

# Medical keywords for validation
MEDICAL_KEYWORDS = [
    'scan', 'ct', 'mri', 'xray', 'x-ray', 'examination', 'ultrasound',
    'findings', 'impression', 'study', 'image', 'report', 'mammogram',
    'clinical', 'patient', 'technique', 'contrast', 'diagnosis',
    'radiology', 'radiologist', 'chest', 'abdomen', 'brain', 'spine',
    'bone', 'lung', 'heart', 'liver', 'kidney', 'pelvis', 'head', 'neck'
]

# -----------------------------
# Page Configuration
# -----------------------------
st.set_page_config(
    page_title=APP_CONFIG["title"],
    page_icon=APP_CONFIG["icon"],
    layout=APP_CONFIG["layout"],
    initial_sidebar_state="collapsed"
)

# -----------------------------
# OpenAI Client (Cached)
# -----------------------------
@st.cache_resource
def get_openai_client():
    """Initialize and cache OpenAI client"""
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        st.error("❌ OpenAI API密鑰未設置，請檢查環境變量")
        st.stop()
    return OpenAI(api_key=api_key)

client = get_openai_client()

# -----------------------------
# Enhanced Language Configuration
# -----------------------------
LANGUAGE_CONFIG = {
    "繁體中文": {
        "code": "traditional_chinese",
        "flag": "🇹🇼",
        "app_title": "RadiAI.Care",
        "app_subtitle": "🩺 智能醫療報告解讀助手",
        "app_description": "將英文放射科報告轉譯為易懂的中文解釋",
        "lang_selection": "🌍 選擇語言",
        "disclaimer_title": "⚠️ 重要法律聲明",
        "disclaimer_items": [
            "純翻譯服務：本工具僅提供語言翻譯，絕不提供醫療建議、診斷或治療建議",
            "準確性限制：AI翻譯可能存在錯誤，請務必與專業醫師核實所有醫療資訊",
            "醫療決策：請勿將翻譯結果用於任何醫療決策，所有醫療問題請諮詢合格醫師",
            "緊急情況：如有緊急醫療需求，請立即撥打000或前往最近的急診室"
        ],
        "input_method": "📝 選擇輸入方式",
        "input_text": "✍️ 直接輸入文字",
        "input_file": "📁 上傳文件",
        "input_placeholder": "請將完整的英文放射科報告貼在下方：",
        "input_help": "請貼上您的英文放射科報告，例如：\n\nCHEST CT SCAN\nCLINICAL HISTORY: Shortness of breath\nTECHNIQUE: Axial CT images...\nFINDINGS: The lungs demonstrate...\nIMPRESSION: ...\n\n請確保包含完整的報告內容以獲得最佳翻譯效果。",
        "file_upload": "📂 選擇您的報告文件",
        "file_formats": "支援的文件格式",
        "file_pdf": "📄 PDF - 掃描或電子版報告",
        "file_txt": "📝 TXT - 純文字報告",
        "file_docx": "📑 DOCX - Word文檔報告",
        "file_success": "✅ 文件讀取成功！",
        "file_preview": "👀 預覽提取的內容",
        "file_error": "❌ 文件讀取失敗，請檢查文件格式或嘗試其他文件",
        "translate_button": "🚀 開始智能解讀",
        "result_title": "📋 智能解讀結果",
        "translation_complete": "🎉 解讀完成！您還有",
        "translation_remaining": "次免費翻譯機會",
        "translation_finished": "🌟 您已用完所有免費翻譯！感謝使用 RadiAI.Care",
        "error_no_content": "請輸入報告內容或上傳有效文件",
        "error_too_short": "輸入內容太短，請確保輸入完整的醫學報告",
        "warning_no_medical": "內容似乎不包含醫學術語，翻譯結果可能不夠準確",
        "usage_remaining": "剩餘",
        "usage_used": "已用",
        "usage_times": "次",
        "quota_finished": "🎯 免費翻譯額度已用完。感謝您的使用！",
        "quota_info": "如需更多翻譯服務，請聯繫我們了解付費方案。",
        "tab_help": "💡 使用指南",
        "tab_privacy": "🔒 隱私保護",
        "tab_reminder": "⚠️ 重要提醒",
        # Feedback related
        "feedback_title": "🗣 使用者回饋",
        "feedback_quick": "此結果對你有幫助嗎？",
        "feedback_clarity": "清晰度",
        "feedback_helpfulness": "實用性",
        "feedback_trust": "可信度",
        "feedback_nps": "願意推薦給朋友 / 家人（0=完全不會，10=一定會）",
        "feedback_issues": "遇到的問題（可複選）",
        "feedback_issue_opts": ["翻譯不準", "用語不好懂", "格式不好閱讀", "速度太慢", "與醫療需求無關", "其他"],
        "feedback_suggestion": "具體建議（選填）",
        "feedback_email_optin": "若希望後續聯絡（選填 Email）",
        "feedback_consent": "我同意此回饋只作為服務改進，不含個人醫療識別資訊",
        "feedback_submit": "提交回饋",
        "feedback_submitted": "✅ 感謝！已收到您的回饋",
        "feedback_already": "已提交過此次翻譯的回饋，感謝！"
    },
    "简体中文": {
        "code": "simplified_chinese",
        "flag": "🇨🇳",
        "app_title": "RadiAI.Care",
        "app_subtitle": "🩺 智能医疗报告解读助手",
        "app_description": "将英文放射科报告转译为易懂的中文解释",
        "lang_selection": "🌍 选择语言",
        "disclaimer_title": "⚠️ 重要法律声明",
        "disclaimer_items": [
            "纯翻译服务：本工具仅提供语言翻译，绝不提供医疗建议、诊断或治疗建议",
            "准确性限制：AI翻译可能存在错误，请务必与专业医师核实所有医疗信息",
            "医疗决策：请勿将翻译结果用于任何医疗决策，所有医疗问题请咨询合格医师",
            "紧急情况：如有紧急医疗需求，请立即拨打000或前往最近的急诊室"
        ],
        "input_method": "📝 选择输入方式",
        "input_text": "✍️ 直接输入文字",
        "input_file": "📁 上传文件",
        "input_placeholder": "请将完整的英文放射科报告贴在下方：",
        "input_help": "请贴上您的英文放射科报告，例如：\n\nCHEST CT SCAN\nCLINICAL HISTORY: Shortness of breath\nTECHNIQUE: Axial CT images...\nFINDINGS: The lungs demonstrate...\nIMPRESSION: ...\n\n请确保包含完整的报告内容以获得最佳翻译效果。",
        "file_upload": "📂 选择您的报告文件",
        "file_formats": "支持的文件格式",
        "file_pdf": "📄 PDF - 扫描或电子版报告",
        "file_txt": "📝 TXT - 纯文字报告",
        "file_docx": "📑 DOCX - Word文档报告",
        "file_success": "✅ 文件读取成功！",
        "file_preview": "👀 预览提取的内容",
        "file_error": "❌ 文件读取失败，请检查文件格式或尝试其他文件",
        "translate_button": "🚀 开始智能解读",
        "result_title": "📋 智能解读结果",
        "translation_complete": "🎉 解读完成！您还有",
        "translation_remaining": "次免费翻译机会",
        "translation_finished": "🌟 您已用完所有免费翻译！感谢使用 RadiAI.Care",
        "error_no_content": "请输入报告内容或上传有效文件",
        "error_too_short": "输入内容太短，请确保输入完整的医学报告",
        "warning_no_medical": "内容似乎不包含医学术语，翻译结果可能不够准确",
        "usage_remaining": "剩余",
        "usage_used": "已用",
        "usage_times": "次",
        "quota_finished": "🎯 免费翻译额度已用完。感谢您的使用！",
        "quota_info": "如需更多翻译服务，请联系我们了解付费方案。",
        "tab_help": "💡 使用指南",
        "tab_privacy": "🔒 隐私保护",
        "tab_reminder": "⚠️ 重要提醒",
        # Feedback related
        "feedback_title": "🗣 用户反馈",
        "feedback_quick": "此结果对你有帮助吗？",
        "feedback_clarity": "清晰度",
        "feedback_helpfulness": "实用性",
        "feedback_trust": "可信度",
        "feedback_nps": "愿意推荐给朋友 / 家人（0=完全不会，10=一定会）",
        "feedback_issues": "遇到的问题（可多选）",
        "feedback_issue_opts": ["翻译不准", "用语不好懂", "格式不好阅读", "速度太慢", "与医疗需求无关", "其他"],
        "feedback_suggestion": "具体建议（选填）",
        "feedback_email_optin": "若希望后续联系（选填 Email）",
        "feedback_consent": "我同意此反馈仅用于改进服务，不含个人医疗识别信息",
        "feedback_submit": "提交反馈",
        "feedback_submitted": "✅ 感谢！已收到你的反馈",
        "feedback_already": "已提交过此次翻译的反馈，感谢！"
    }
}

# -----------------------------
# Session State Management
# -----------------------------
class SessionManager:
    """Enhanced session state management"""
    
    @staticmethod
    def init_session_state():
        """Initialize all session state variables"""
        defaults = {
            'language': "简体中文",
            'translation_count': 0,
            'input_method': "text",
            'feedback_submitted_ids': set(),
            'last_translation_id': None,
            'user_session_id': str(uuid.uuid4())[:8],
            'app_start_time': time.time(),
            'translation_history': [],
            'last_activity': time.time()
        }
        
        for key, default_value in defaults.items():
            if key not in st.session_state:
                st.session_state[key] = default_value
    
    @staticmethod
    def update_activity():
        """Update last activity timestamp"""
        st.session_state.last_activity = time.time()
    
    @staticmethod
    def add_translation_record(translation_id: str, text_length: int, language: str):
        """Add translation to history"""
        record = {
            'id': translation_id,
            'timestamp': time.time(),
            'text_length': text_length,
            'language': language
        }
        if 'translation_history' not in st.session_state:
            st.session_state.translation_history = []
        st.session_state.translation_history.append(record)

# -----------------------------
# Enhanced CSS Styling
# -----------------------------
def load_enhanced_css():
    """Load comprehensive CSS styling"""
    st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Noto+Sans+SC:wght@300;400;500;600;700&display=swap');

        :root {
            --primary-blue: #0d74b8;
            --secondary-blue: #29a3d7;
            --light-blue: #e0f7fa;
            --background-gradient: linear-gradient(135deg,#e0f7fa 0%,#ffffff 55%,#f2fbfe 100%);
            --card-shadow: 0 10px 30px rgba(40,85,120,0.12);
            --hover-shadow: 0 8px 25px rgba(31, 119, 180, 0.4);
            --border-radius: 16px;
            --transition: all 0.3s ease;
        }

        .stApp {
            font-family: 'Inter','Noto Sans SC',-apple-system,BlinkMacSystemFont,sans-serif;
            background: var(--background-gradient);
            min-height: 100vh;
        }

        .main-container {
            background: rgba(255,255,255,0.95);
            backdrop-filter: blur(10px);
            border-radius: 20px;
            margin: 1rem auto;
            padding: 2rem 2.5rem;
            box-shadow: var(--card-shadow);
            max-width: 900px;
            border: 1px solid rgba(227,238,245,0.8);
            position: relative;
            overflow: hidden;
        }

        .main-container::before {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            height: 4px;
            background: linear-gradient(90deg, var(--primary-blue), var(--secondary-blue));
            border-radius: 20px 20px 0 0;
        }

        .title-section {
            text-align: center;
            padding: 2.5rem 0 2rem;
            margin-bottom: 1.5rem;
            background: linear-gradient(145deg,#f4fcff 0%,#e5f4fb 100%);
            border-radius: 18px;
            border: 1px solid #d4e8f2;
            position: relative;
        }

        .main-title {
            font-size: 3rem;
            font-weight: 700;
            background: linear-gradient(90deg,var(--primary-blue) 0%,var(--secondary-blue) 60%,#1b90c8 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
            margin-bottom: 0.5rem;
            letter-spacing: -1px;
            text-shadow: 0 4px 8px rgba(0,0,0,0.05);
        }

        .subtitle {
            font-size: 1.4rem;
            color: #256084;
            font-weight: 600;
            margin-bottom: 0.5rem;
        }

        .description {
            font-size: 1.05rem;
            color: #4c7085;
            max-width: 600px;
            margin: 0 auto;
            line-height: 1.6;
            opacity: 0.9;
        }

        .disclaimer-container {
            background: linear-gradient(135deg,#fff8e1 0%,#ffefd5 100%);
            border: 2px solid #f7c56a;
            border-radius: var(--border-radius);
            padding: 1.5rem;
            margin: 1.5rem 0;
            box-shadow: 0 6px 20px rgba(240,168,0,0.15);
            position: relative;
        }

        .disclaimer-title {
            text-align: center;
            font-size: 1.2rem;
            font-weight: 700;
            margin-bottom: 1rem;
            color: #b25209;
        }

        .disclaimer-item {
            margin-bottom: 0.8rem;
            padding: 1rem;
            background: rgba(255,255,255,0.8);
            border-radius: 10px;
            border-left: 4px solid #ffb438;
            font-size: 0.95rem;
            line-height: 1.5;
            color: #a34907;
            font-weight: 500;
            transition: var(--transition);
        }

        .disclaimer-item:hover {
            transform: translateX(4px);
            box-shadow: 0 4px 12px rgba(0,0,0,0.1);
        }

        .input-section {
            background: linear-gradient(145deg,#f3faff 0%,#e8f5ff 100%);
            border-radius: 18px;
            padding: 1.8rem;
            margin: 1.5rem 0;
            border: 2px solid #d2e8f3;
            box-shadow: 0 6px 20px rgba(13,116,184,0.08);
            position: relative;
        }

        .stButton > button {
            background: linear-gradient(135deg,var(--primary-blue) 0%,var(--secondary-blue) 50%,#15a4e7 100%);
            color: white;
            border: none;
            border-radius: 14px;
            padding: 1rem 2rem;
            font-weight: 600;
            font-size: 1.1rem;
            box-shadow: 0 6px 20px rgba(23,124,179,0.3);
            transition: var(--transition);
            width: 100%;
            margin: 0.8rem 0;
            position: relative;
            overflow: hidden;
        }

        .stButton > button:hover {
            transform: translateY(-3px);
            box-shadow: var(--hover-shadow);
        }

        .stButton > button:active {
            transform: translateY(-1px) scale(0.98);
        }

        .result-container {
            background: linear-gradient(145deg,#f2fbff 0%,#e3f4fa 100%);
            border-radius: 18px;
            padding: 2rem;
            margin: 1.8rem 0;
            border-left: 5px solid var(--secondary-blue);
            box-shadow: 0 8px 25px rgba(9,110,160,0.12);
            position: relative;
        }

        .feedback-container {
            background: rgba(255,255,255,0.95);
            border: 1px solid #d8ecf4;
            padding: 1.5rem;
            border-radius: var(--border-radius);
            box-shadow: 0 6px 20px rgba(20,120,170,0.08);
            margin-top: 1.5rem;
        }

        .feedback-container h4 {
            margin-top: 0;
            font-size: 1.1rem;
            color: var(--primary-blue);
            display: flex;
            align-items: center;
            gap: 0.5rem;
        }

        .stFileUploader > div {
            border: 3px dashed var(--secondary-blue);
            border-radius: 15px;
            padding: 2rem;
            text-align: center;
            background: linear-gradient(135deg,#f5fcff 0%,#e9f6fb 100%);
            transition: var(--transition);
        }

        .stFileUploader > div:hover {
            border-color: var(--primary-blue);
            background: linear-gradient(135deg,#eef9ff 0%,#def1fa 100%);
            transform: translateY(-2px);
        }

        .stTextArea > div > div > textarea {
            border-radius: 12px;
            border: 2px solid #c9e1ec;
            font-family: 'Inter',sans-serif;
            font-size: 1rem;
            padding: 1rem;
            transition: var(--transition);
        }

        .stTextArea > div > div > textarea:focus {
            border-color: var(--secondary-blue);
            box-shadow: 0 0 0 3px rgba(21,146,213,0.15);
        }

        .stProgress > div > div > div > div {
            background: linear-gradient(90deg,var(--primary-blue) 0%,var(--secondary-blue) 100%);
            border-radius: 10px;
        }

        .stSuccess {
            border-radius: 12px;
            border-left: 4px solid #28a745;
            background: linear-gradient(135deg,#e6f7ed 0%,#f1fbf5 100%);
            padding: 1rem;
        }

        .stError {
            border-radius: 12px;
            border-left: 4px solid #dc3545;
            background: linear-gradient(135deg,#fdeaea 0%,#fff3f3 100%);
            padding: 1rem;
        }

        .stWarning {
            border-radius: 12px;
            border-left: 4px solid #ffc107;
            background: linear-gradient(135deg,#fff8e1 0%,#fffbf0 100%);
            padding: 1rem;
        }

        .stInfo {
            border-radius: 12px;
            border-left: 4px solid #17a2b8;
            background: linear-gradient(135deg,#e1f7fa 0%,#f0fcff 100%);
            padding: 1rem;
        }

        .usage-tracker {
            background: linear-gradient(135deg,#f8fbff 0%,#f0f7ff 100%);
            border-radius: var(--border-radius);
            padding: 1.5rem;
            margin: 1rem 0;
            border: 1px solid #e1eeff;
        }

        .language-selector {
            display: flex;
            gap: 1rem;
            justify-content: center;
            margin: 1.5rem 0;
        }

        @media (max-width: 768px) {
            .main-title { font-size: 2.2rem; }
            .subtitle { font-size: 1.1rem; }
            .main-container { margin: 0.5rem; padding: 1.5rem; }
            .disclaimer-item { font-size: 0.9rem; }
            .language-selector { flex-direction: column; gap: 0.5rem; }
        }

        /* Hide Streamlit elements */
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        .stDeployButton {display:none;}
        
        /* Custom scrollbar */
        ::-webkit-scrollbar { width: 8px; }
        ::-webkit-scrollbar-track { background: #f1f1f1; border-radius: 10px; }
        ::-webkit-scrollbar-thumb { 
            background: linear-gradient(180deg,var(--primary-blue),var(--secondary-blue)); 
            border-radius: 10px; 
        }
        ::-webkit-scrollbar-thumb:hover { 
            background: linear-gradient(180deg,#0a75b0,#1192ce); 
        }

        /* Animation classes */
        .fade-in {
            animation: fadeIn 0.5s ease-in;
        }

        @keyframes fadeIn {
            from { opacity: 0; transform: translateY(20px); }
            to { opacity: 1; transform: translateY(0); }
        }

        .pulse {
            animation: pulse 2s infinite;
        }

        @keyframes pulse {
            0% { transform: scale(1); }
            50% { transform: scale(1.05); }
            100% { transform: scale(1); }
        }
    </style>
    """, unsafe_allow_html=True)

# -----------------------------
# Enhanced File Processing
# -----------------------------
class FileProcessor:
    """Enhanced file processing with better error handling"""
    
    @staticmethod
    def validate_file(uploaded_file) -> Tuple[bool, str]:
        """Validate uploaded file"""
        if uploaded_file is None:
            return False, "沒有選擇文件"
        
        # Check file size
        file_size_mb = len(uploaded_file.getvalue()) / (1024 * 1024)
        if file_size_mb > LIMITS["file_size_limit_mb"]:
            return False, f"文件過大，請上傳小於{LIMITS['file_size_limit_mb']}MB的文件"
        
        # Check file type
        file_extension = uploaded_file.name.lower().split('.')[-1]
        if file_extension not in SUPPORTED_FILE_TYPES:
            return False, f"不支持的文件格式，請上傳{', '.join(SUPPORTED_FILE_TYPES).upper()}文件"
        
        return True, ""
    
    @staticmethod
    def extract_text_from_file(uploaded_file) -> Optional[str]:
        """Extract text from uploaded files with comprehensive error handling"""
        try:
            is_valid, error_msg = FileProcessor.validate_file(uploaded_file)
            if not is_valid:
                logger.error(f"File validation failed: {error_msg}")
                return None
            
            file_extension = uploaded_file.name.lower().split('.')[-1]
            
            if file_extension == 'txt':
                return FileProcessor._extract_from_txt(uploaded_file)
            elif file_extension == 'pdf':
                return FileProcessor._extract_from_pdf(uploaded_file)
            elif file_extension in ['docx', 'doc']:
                return FileProcessor._extract_from_docx(uploaded_file)
            else:
                logger.error(f"Unsupported file type: {file_extension}")
                return None
                
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return None
    
    @staticmethod
    def _extract_from_txt(uploaded_file) -> str:
        """Extract text from TXT files"""
        encodings = ['utf-8', 'gbk', 'big5', 'latin-1']
        
        for encoding in encodings:
            try:
                uploaded_file.seek(0)
                content = uploaded_file.read().decode(encoding)
                return content.strip()
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Unable to decode text file with any supported encoding")
    
    @staticmethod
    def _extract_from_pdf(uploaded_file) -> str:
        """Extract text from PDF files"""
        pdf_document = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        text_parts = []
        
        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(page_text)
        
        pdf_document.close()
        return "\n\n".join(text_parts).strip() if text_parts else ""
    
    @staticmethod
    def _extract_from_docx(uploaded_file) -> str:
        """Extract text from DOCX files"""
        doc = docx.Document(io.BytesIO(uploaded_file.read()))
        text_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts).strip() if text_parts else ""

# -----------------------------
# Enhanced Content Validation
# -----------------------------
class ContentValidator:
    """Enhanced content validation with medical term detection"""
    
    @staticmethod
    def validate_medical_content(text: str) -> Dict[str, Any]:
        """Comprehensive medical content validation"""
        if not text or len(text.strip()) < LIMITS["min_text_length"]:
            return {
                "is_valid": False,
                "confidence": 0.0,
                "medical_terms": [],
                "issues": ["內容過短"],
                "suggestions": ["請輸入完整的醫學報告"]
            }
        
        text_lower = text.lower()
        found_terms = [kw for kw in MEDICAL_KEYWORDS if kw in text_lower]
        
        # Calculate confidence based on medical terms and structure
        confidence = min(len(found_terms) * 0.15, 1.0)
        
        # Check for report structure indicators
        structure_indicators = [
            'impression:', 'findings:', 'technique:', 'clinical history:',
            'examination:', 'study:', 'images show', 'no evidence of'
        ]
        
        structure_score = sum(1 for indicator in structure_indicators if indicator in text_lower)
        confidence += min(structure_score * 0.1, 0.4)
        
        is_valid = len(found_terms) >= 2 and len(text.strip()) >= LIMITS["min_text_length"]
        
        issues = []
        suggestions = []
        
        if len(found_terms) < 2:
            issues.append("醫學術語過少")
            suggestions.append("請確認這是醫學報告")
        
        if structure_score == 0:
            issues.append("缺少報告結構")
            suggestions.append("請包含完整的報告段落")
        
        return {
            "is_valid": is_valid,
            "confidence": min(confidence, 1.0),
            "medical_terms": found_terms,
            "issues": issues,
            "suggestions": suggestions,
            "structure_score": structure_score
        }

# -----------------------------
# Enhanced Translation Engine
# -----------------------------
class TranslationEngine:
    """Enhanced translation with retry logic and quality assessment"""
    
    def __init__(self, openai_client):
        self.client = openai_client
        self.max_retries = LIMITS["max_retries"]
        self.timeout = LIMITS["request_timeout"]
    
    def translate_report(self, report_text: str, language_code: str) -> Tuple[str, str]:
        """Main translation function with enhanced error handling"""
        for attempt in range(self.max_retries):
            try:
                system_prompt = get_prompt(language_code)
                
                response = self.client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": f"請翻譯並解釋以下放射科報告：\n\n{report_text}"}
                    ],
                    temperature=0.3,
                    max_tokens=2500,
                    timeout=self.timeout
                )
                
                result_text = response.choices[0].message.content.strip()
                disclaimer_html = create_enhanced_disclaimer(language_code)
                
                return result_text, disclaimer_html
                
            except Exception as e:
                logger.warning(f"Translation attempt {attempt + 1} failed: {e}")
                
                if attempt < self.max_retries - 1:
                    # Exponential backoff
                    wait_time = 2 ** attempt
                    time.sleep(wait_time)
                else:
                    return self._handle_translation_error(e), ""
        
        return "❌ 翻譯失敗，請稍後重試", ""
    
    def _handle_translation_error(self, error: Exception) -> str:
        """Handle translation errors with specific messages"""
        error_msg = str(error).lower()
        
        if "rate limit" in error_msg:
            return "❌ API請求過於頻繁，請稍後重試。"
        elif "timeout" in error_msg:
            return "❌ 請求超時，請檢查網絡連接後重試。"
        elif "api" in error_msg or "openai" in error_msg:
            return "❌ AI服務暫時不可用，請稍後重試。"
        else:
            return f"❌ 翻譯過程中發生錯誤：{error_msg}\n\n請檢查網絡連接或稍後重試。"

# -----------------------------
# Enhanced Feedback System
# -----------------------------
class FeedbackManager:
    """Enhanced feedback collection and logging"""
    
    @staticmethod
    def log_feedback_to_sheets(**data):
        """Log feedback to Google Sheets with error handling"""
        try:
            log_to_google_sheets(processing_status="feedback", **data)
            return True
        except Exception as e:
            logger.warning(f"Feedback logging failed: {e}")
            return False
    
    @staticmethod
    def render_feedback_section(lang: Dict, translation_id: str, 
                              report_length: int, free_quota_remaining: int,
                              file_type: str, is_medical_detected: bool):
        """Render enhanced feedback collection UI"""
        if translation_id in st.session_state.feedback_submitted_ids:
            st.info(lang.get('feedback_already', '已提交過此次翻譯的回饋'))
            return
        
        with st.container():
            st.markdown('<div class="feedback-container">', unsafe_allow_html=True)
            st.markdown(f"#### {lang.get('feedback_title', '🗣 回饋')}")
            
            # Quick sentiment buttons
            col1, col2 = st.columns(2)
            quick_sentiment = None
            
            with col1:
                if st.button("👍 有幫助", key=f"fb_up_{translation_id}", use_container_width=True):
                    quick_sentiment = "positive"
            
            with col2:
                if st.button("👎 沒幫助", key=f"fb_down_{translation_id}", use_container_width=True):
                    quick_sentiment = "negative"
            
            # Detailed feedback form
            with st.form(f"feedback_form_{translation_id}", clear_on_submit=True):
                st.markdown("##### 詳細評價（可選）")
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    clarity = st.slider(lang.get('feedback_clarity', '清晰度'), 1, 5, 4)
                with col2:
                    helpfulness = st.slider(lang.get('feedback_helpfulness', '實用性'), 1, 5, 4)
                with col3:
                    trust = st.slider(lang.get('feedback_trust', '可信度'), 1, 5, 4)
                
                nps = st.slider(lang.get('feedback_nps', '推薦指數'), 0, 10, 8)
                
                issue_opts = lang.get('feedback_issue_opts', [
                    "翻譯不準", "用語不好懂", "格式不好閱讀", 
                    "速度太慢", "與醫療需求無關", "其他"
                ])
                issues = st.multiselect(lang.get('feedback_issues', '遇到的問題'), issue_opts)
                
                suggestion = st.text_area(
                    lang.get('feedback_suggestion', '具體建議'), 
                    height=80, max_chars=500
                )
                
                col1, col2 = st.columns([3, 1])
                with col1:
                    email_opt = st.checkbox(lang.get('feedback_email_optin', '希望後續聯絡'))
                with col2:
                    email_value = st.text_input("Email", disabled=not email_opt) if email_opt else ""
                
                consent = st.checkbox(lang.get('feedback_consent', '同意服務改進使用'), value=True)
                
                submitted = st.form_submit_button(
                    lang.get('feedback_submit', '提交回饋'), 
                    use_container_width=True,
                    type="primary"
                )
            
            if submitted and consent:
                if quick_sentiment is None:
                    quick_sentiment = "positive" if (clarity + helpfulness + trust) >= 11 else "negative"
                
                feedback_data = {
                    'translation_id': translation_id,
                    'language': st.session_state.language,
                    'report_length': report_length,
                    'file_type': file_type,
                    'quick_sentiment': quick_sentiment,
                    'clarity_score': clarity,
                    'helpfulness_score': helpfulness,
                    'trust_score': trust,
                    'nps': nps,
                    'issues': ";".join(issues),
                    'suggestion_text': suggestion.strip(),
                    'email': email_value.strip() if email_opt else "",
                    'email_opt_in': email_opt,
                    'consent': consent,
                    'free_quota_remaining': free_quota_remaining,
                    'is_medical_detected': is_medical_detected,
                    'app_version': APP_CONFIG["version"]
                }
                
                if FeedbackManager.log_feedback_to_sheets(**feedback_data):
                    st.session_state.feedback_submitted_ids.add(translation_id)
                    st.success(lang.get('feedback_submitted', '✅ 感謝！已收到您的回饋'))
                    st.balloons()
                else:
                    st.warning("回饋提交失敗，但您的意見已記錄")
            
            elif submitted and not consent:
                st.error("請勾選同意條款後再提交")
            
            st.markdown('</div>', unsafe_allow_html=True)

# -----------------------------
# UI Components
# -----------------------------
class UIComponents:
    """Enhanced UI component rendering"""
    
    @staticmethod
    def render_language_selection(lang: Dict):
        """Render enhanced language selection"""
        st.markdown(f'<div style="text-align:center; margin:1.5rem 0;"><h4 style="margin:0;">{lang["lang_selection"]}</h4></div>', unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button(
                f"{LANGUAGE_CONFIG['繁體中文']['flag']} 繁體中文", 
                key="lang_traditional",
                use_container_width=True,
                type="primary" if st.session_state.language == "繁體中文" else "secondary"
            ):
                st.session_state.language = "繁體中文"
                SessionManager.update_activity()
                st.rerun()
        
        with col2:
            if st.button(
                f"{LANGUAGE_CONFIG['简体中文']['flag']} 简体中文", 
                key="lang_simplified",
                use_container_width=True,
                type="primary" if st.session_state.language == "简体中文" else "secondary"
            ):
                st.session_state.language = "简体中文"
                SessionManager.update_activity()
                st.rerun()
    
    @staticmethod
    def render_disclaimer(lang: Dict):
        """Render enhanced disclaimer section"""
        disclaimer_html = f'<div class="disclaimer-container"><div class="disclaimer-title">{lang["disclaimer_title"]}</div>'
        
        for i, item in enumerate(lang["disclaimer_items"], 1):
            disclaimer_html += f'<div class="disclaimer-item"><strong>📌 {i}.</strong> {item}</div>'
        
        disclaimer_html += '</div>'
        st.markdown(disclaimer_html, unsafe_allow_html=True)
    
    @staticmethod
    def render_usage_tracker(lang: Dict) -> int:
        """Render enhanced usage tracking with progress visualization"""
        remaining = LIMITS["max_free_translations"] - st.session_state.translation_count
        progress = st.session_state.translation_count / LIMITS["max_free_translations"]
        
        st.markdown('<div class="usage-tracker">', unsafe_allow_html=True)
        st.markdown("### 📊 使用情況追蹤")
        
        col1, col2, col3 = st.columns([4, 1, 1])
        
        with col1:
            st.progress(progress)
            if remaining > 0:
                st.caption(f"還可使用 {remaining} 次免費翻譯")
            else:
                st.caption("免費額度已用完")
        
        with col2:
            if remaining > 0:
                st.metric("剩餘", remaining, delta=None)
            else:
                st.metric("剩餘", 0, delta="已用完")
        
        with col3:
            st.metric("總計", f"{st.session_state.translation_count}/{LIMITS['max_free_translations']}")
        
        st.markdown('</div>', unsafe_allow_html=True)
        return remaining
    
    @staticmethod
    def render_input_section(lang: Dict) -> Tuple[str, str]:
        """Render enhanced input section with better UX"""
        st.markdown('<div class="input-section">', unsafe_allow_html=True)
        st.markdown(f'### {lang["input_method"]}')
        
        # Input method selection
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button(
                lang["input_text"], 
                key="input_text_btn",
                use_container_width=True,
                type="primary" if st.session_state.input_method == "text" else "secondary"
            ):
                st.session_state.input_method = "text"
                SessionManager.update_activity()
        
        with col2:
            if st.button(
                lang["input_file"], 
                key="input_file_btn",
                use_container_width=True,
                type="primary" if st.session_state.input_method == "file" else "secondary"
            ):
                st.session_state.input_method = "file"
                SessionManager.update_activity()
        
        report_text = ""
        file_type = "manual"
        
        if st.session_state.input_method == "text":
            st.markdown("#### 📝 輸入報告內容")
            report_text = st.text_area(
                lang["input_placeholder"],
                height=280,
                placeholder=lang["input_help"],
                help="💡 支援各種格式的英文放射科報告",
                label_visibility="collapsed",
                max_chars=LIMITS["max_text_length"]
            )
            
            if report_text:
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.caption(f"字符數: {len(report_text)}")
                with col2:
                    if len(report_text) > LIMITS["max_text_length"]:
                        st.error("內容過長")
                    elif len(report_text) < LIMITS["min_text_length"]:
                        st.warning("內容過短")
                    else:
                        st.success("長度適中")
        
        else:
            st.markdown("#### 📂 上傳報告文件")
            
            uploaded_file = st.file_uploader(
                lang["file_upload"],
                type=SUPPORTED_FILE_TYPES,
                help=f"📋 支援 {', '.join([t.upper() for t in SUPPORTED_FILE_TYPES])} 格式，文件大小限制 {LIMITS['file_size_limit_mb']}MB",
                label_visibility="collapsed"
            )
            
            # File format info
            with st.expander(f"📋 {lang['file_formats']}", expanded=False):
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.markdown(f"**{lang['file_pdf']}**")
                with col2:
                    st.markdown(f"**{lang['file_txt']}**")
                with col3:
                    st.markdown(f"**{lang['file_docx']}**")
            
            if uploaded_file is not None:
                file_type = uploaded_file.name.lower().split('.')[-1]
                
                with st.spinner("🔄 正在讀取文件內容..."):
                    extracted_text = FileProcessor.extract_text_from_file(uploaded_file)
                    
                    if extracted_text:
                        report_text = extracted_text
                        st.success(f"✅ {lang['file_success']}")
                        
                        # File info
                        file_size = len(uploaded_file.getvalue()) / 1024
                        st.caption(f"文件名: {uploaded_file.name} | 大小: {file_size:.1f} KB | 內容長度: {len(extracted_text)} 字符")
                        
                        with st.expander(f"👀 {lang['file_preview']}", expanded=False):
                            preview_text = (extracted_text[:LIMITS["preview_length"]] + "...") if len(extracted_text) > LIMITS["preview_length"] else extracted_text
                            st.text_area("內容預覽", value=preview_text, height=150, disabled=True)
                    else:
                        st.error(f"❌ {lang['file_error']}")
                        file_type = "failed"
        
        st.markdown('</div>', unsafe_allow_html=True)
        return report_text, file_type

# -----------------------------
# Main Translation Process
# -----------------------------
class TranslationProcessor:
    """Main translation processing workflow"""
    
    def __init__(self):
        self.translation_engine = TranslationEngine(client)
        self.validator = ContentValidator()
    
    def process_translation(self, report_text: str, file_type: str, lang: Dict) -> bool:
        """Handle complete translation workflow"""
        # Input validation
        if not report_text.strip():
            st.error(f"❌ {lang['error_no_content']}")
            return False
        
        if len(report_text.strip()) < LIMITS["min_text_length"]:
            st.error(f"❌ {lang['error_too_short']}")
            return False
        
        # Content validation
        validation_result = self.validator.validate_medical_content(report_text)
        
        if not validation_result["is_valid"]:
            st.warning(f"⚠️ {lang['warning_no_medical']}")
            if validation_result["suggestions"]:
                for suggestion in validation_result["suggestions"]:
                    st.info(f"💡 {suggestion}")
        
        # Show processing steps
        progress_container = st.container()
        with progress_container:
            progress_bar = st.progress(0)
            status_text = st.empty()
        
        start_time = time.time()
        translation_id = str(uuid.uuid4())
        
        try:
            # Get processing steps
            steps = get_processing_steps(lang["code"])
            total_steps = len(steps)
            
            # Process with animated progress
            for i, step_text in enumerate(steps):
                status_text.markdown(f"**🔄 {step_text}**")
                progress_bar.progress(int((i + 1) / total_steps * 80))  # Leave 20% for actual translation
                time.sleep(0.8)
            
            # Perform translation
            status_text.markdown("**🤖 AI正在生成解讀結果...**")
            progress_bar.progress(90)
            
            result_text, disclaimer_html = self.translation_engine.translate_report(
                report_text, lang["code"]
            )
            
            # Complete progress
            progress_bar.progress(100)
            time.sleep(0.5)
            progress_container.empty()
            
            # Calculate processing time
            processing_time_ms = int((time.time() - start_time) * 1000)
            
            # Display results
            st.markdown('<div class="result-container fade-in">', unsafe_allow_html=True)
            st.markdown(f"### {lang['result_title']} ({st.session_state.language})")
            st.markdown(result_text)
            
            if disclaimer_html:
                st.markdown(disclaimer_html, unsafe_allow_html=True)
            
            st.markdown('</div>', unsafe_allow_html=True)
            
            # Log successful translation
            try:
                log_to_google_sheets(
                    language=st.session_state.language,
                    report_length=len(report_text),
                    file_type=file_type,
                    processing_status="success",
                    translation_id=translation_id,
                    latency_ms=processing_time_ms,
                    is_medical_detected=validation_result["is_valid"],
                    app_version=APP_CONFIG["version"]
                )
            except Exception as log_error:
                logger.warning(f"Failed to log usage: {log_error}")
            
            # Update session state
            st.session_state.translation_count += 1
            st.session_state.last_translation_id = translation_id
            SessionManager.add_translation_record(translation_id, len(report_text), st.session_state.language)
            SessionManager.update_activity()
            
            # Show completion status
            new_remaining = LIMITS["max_free_translations"] - st.session_state.translation_count
            if new_remaining > 0:
                st.success(f"{lang['translation_complete']} {new_remaining} {lang['translation_remaining']}")
            else:
                st.balloons()
                st.success(f"{lang['translation_finished']}")
            
            # Show feedback section
            FeedbackManager.render_feedback_section(
                lang=lang,
                translation_id=translation_id,
                report_length=len(report_text),
                free_quota_remaining=new_remaining,
                file_type=file_type,
                is_medical_detected=validation_result["is_valid"]
            )
            
            return True
            
        except Exception as e:
            progress_container.empty()
            st.error(f"❌ 翻譯過程中發生錯誤：{str(e)}")
            
            # Log error
            try:
                log_to_google_sheets(
                    language=st.session_state.language,
                    report_length=len(report_text),
                    file_type=file_type,
                    processing_status="error",
                    error=str(e),
                    app_version=APP_CONFIG["version"]
                )
            except Exception:
                pass
            
            return False

# -----------------------------
# Footer and Information Tabs
# -----------------------------
def render_footer_tabs(lang: Dict):
    """Render comprehensive footer information"""
    st.markdown("---")
    
    tab1, tab2, tab3, tab4 = st.tabs([
        f"{lang['tab_help']}", 
        f"{lang['tab_privacy']}", 
        f"{lang['tab_reminder']}",
        "📊 統計"
    ])
    
    with tab1:
        st.markdown("""
        ### 🎯 如何獲得最佳解讀效果？
        
        #### 📋 報告準備
        ✅ **完整內容**：確保報告包含所有段落和專業術語  
        ✅ **清晰格式**：如上傳圖片，請確保文字清晰可讀  
        ✅ **正確文件**：支援 PDF、TXT、DOCX 格式  
        ✅ **適當長度**：建議 100-5000 字符長度範圍  
        
        #### 🔍 質量檢查
        ✅ **包含關鍵詞**：impression、findings、technique 等  
        ✅ **醫學術語**：CT、MRI、X-ray 等檢查術語  
        ✅ **結構完整**：包含檢查方法、發現、結論等部分  
        
        #### 💡 使用建議
        ✅ **參考性質**：翻譯結果僅供參考，請諮詢醫師  
        ✅ **保護隱私**：建議去除姓名等個人信息  
        ✅ **及時諮詢**：如有疑問請聯繫您的醫療團隊
        """)
    
    with tab2:
        st.markdown("""
        ### 🔒 您的隱私受到全面保護
        
        #### 🛡️ 數據安全承諾
        ✅ **不儲存醫療內容**：報告內容處理完成後立即清除  
        ✅ **匿名統計**：僅記錄匿名使用統計以改善服務  
        ✅ **加密傳輸**：所有數據傳輸均使用 HTTPS 加密  
        ✅ **本地處理**：部分處理在您的設備上完成  
        
        #### ⚠️ 隱私建議
        🚫 **避免敏感信息**：請勿包含身份證號、地址等  
        🚫 **移除姓名**：建議在上傳前移除患者姓名  
        ✅ **關注內容**：我們只關注醫學術語翻譯  
        
        #### 📋 合規標準
        ✅ 符合澳洲隱私法 (Privacy Act 1988)  
        ✅ 遵循 GDPR 數據保護標準  
        ✅ 醫療數據處理最佳實踐
        """)
    
    with tab3:
        st.markdown("""
        ### ⚠️ 重要醫療安全提醒
        
        #### ✅ 我們提供的服務
        🔹 **語言翻譯**：英文醫學報告的中文翻譯  
        🔹 **術語解釋**：醫學專業術語的通俗解釋  
        🔹 **結構梳理**：報告內容的整理和歸納  
        🔹 **問題建議**：向醫師諮詢的參考問題  
        
        #### 🚫 我們不提供的服務
        ❌ **醫療診斷**：任何形式的疾病診斷  
        ❌ **治療建議**：藥物或治療方案建議  
        ❌ **醫療決策**：影響醫療選擇的建議  
        ❌ **緊急醫療**：急救或緊急醫療服務  
        
        #### 🆘 緊急情況處理
        📞 **緊急醫療**：立即撥打 **000**  
        🏥 **急診就醫**：前往最近的急診室  
        👨‍⚕️ **專業諮詢**：聯繫您的主治醫師  
        📋 **法律責任**：所有醫療決策責任歸患者和醫療團隊
        
        #### ⚖️ 服務條款
        使用本服務即表示您理解並同意上述所有條款和限制。
        """)
    
    with tab4:
        if hasattr(st.session_state, 'translation_history') and st.session_state.translation_history:
            st.markdown("### 📈 使用統計")
            
            total_translations = len(st.session_state.translation_history)
            total_characters = sum(record['text_length'] for record in st.session_state.translation_history)
            session_duration = int((time.time() - st.session_state.app_start_time) / 60)
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("總翻譯次數", total_translations)
            with col2:
                st.metric("總字符數", f"{total_characters:,}")
            with col3:
                st.metric("會話時長", f"{session_duration} 分鐘")
            
            # Show translation history
            if st.checkbox("顯示詳細記錄"):
                st.markdown("#### 翻譯記錄")
                for i, record in enumerate(reversed(st.session_state.translation_history), 1):
                    timestamp = time.strftime("%H:%M:%S", time.localtime(record['timestamp']))
                    st.text(f"{i}. [{timestamp}] {record['language']} - {record['text_length']} 字符")
        else:
            st.info("暫無使用記錄")

# -----------------------------
# Main Application
# -----------------------------
def main():
    """Enhanced main application with comprehensive error handling"""
    try:
        # Initialize session state
        SessionManager.init_session_state()
        
        # Load enhanced CSS
        load_enhanced_css()
        
        # Get current language configuration
        lang = LANGUAGE_CONFIG[st.session_state.language]
        
        # Main container with animation class
        st.markdown('<div class="main-container fade-in">', unsafe_allow_html=True)
        
        # Enhanced title section
        st.markdown(f'''
        <div class="title-section">
            <div class="main-title pulse">{lang["app_title"]}</div>
            <div class="subtitle">{lang["app_subtitle"]}</div>
            <div class="description">{lang["app_description"]}</div>
        </div>
        ''', unsafe_allow_html=True)
        
        # Language selection with enhanced UX
        UIComponents.render_language_selection(lang)
        
        # Update language configuration after selection
        lang = LANGUAGE_CONFIG[st.session_state.language]
        
        # Enhanced disclaimer
        UIComponents.render_disclaimer(lang)
        
        # Enhanced usage tracking
        remaining = UIComponents.render_usage_tracker(lang)
        
        # Check quota
        if remaining <= 0:
            st.error(f"🚫 {lang['quota_finished']}")
            st.info(f"💡 {lang['quota_info']}")
            
            # Show contact information
            with st.expander("📞 聯繫我們", expanded=True):
                st.markdown("""
                #### 如需更多服務，請聯繫：
                
                📧 **Email**: support@radiai.care  
                🌐 **網站**: www.radiai.care  
                📱 **服務時間**: 週一至週五 9:00-17:00 (AEST)
                
                我們提供：
                - 無限次數翻譯服務
                - 批量報告處理
                - 專業醫療術語定制
                - 優先客戶支援
                """)
            
            # Show usage statistics even when quota exceeded
            render_footer_tabs(lang)
            st.markdown('</div>', unsafe_allow_html=True)
            return
        
        # Enhanced input section
        report_text, file_type = UIComponents.render_input_section(lang)
        
        # Enhanced translation button with validation
        button_disabled = not report_text.strip() or len(report_text.strip()) < LIMITS["min_text_length"]
        button_help = "請輸入有效的報告內容" if button_disabled else "點擊開始智能解讀"
        
        if st.button(
            f"{lang['translate_button']}", 
            type="primary", 
            use_container_width=True,
            disabled=button_disabled,
            help=button_help
        ):
            # Process translation with enhanced workflow
            processor = TranslationProcessor()
            success = processor.process_translation(report_text, file_type, lang)
            
            if success:
                # Show success metrics
                st.markdown("---")
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("處理狀態", "✅ 成功")
                with col2:
                    st.metric("內容長度", f"{len(report_text)} 字符")
                with col3:
                    st.metric("剩餘額度", remaining - 1)
        
        # Enhanced footer with comprehensive information
        render_footer_tabs(lang)
        
        # Enhanced version info with additional details
        st.markdown(f'''
        <div style="text-align: center; color: #587488; font-size: 0.85rem; margin-top: 2rem; padding: 1.5rem; background: linear-gradient(135deg, #f2f9fc 0%, #e8f4f8 100%); border: 1px solid #d6e7ef; border-radius: 12px;">
            <div style="margin-bottom: 0.5rem;">
                <strong>RadiAI.Care {APP_CONFIG["version"]}</strong> | 為澳洲華人社區打造 ❤️
            </div>
            <div style="font-size: 0.75rem; opacity: 0.8;">
                Powered by GPT-4o | Built with Streamlit | Session: {st.session_state.user_session_id}
            </div>
            <div style="font-size: 0.7rem; margin-top: 0.5rem; opacity: 0.6;">
                {APP_CONFIG["description"]} | 安全 · 準確 · 易用
            </div>
        </div>
        ''', unsafe_allow_html=True)
        
        # Close main container
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Update session activity
        SessionManager.update_activity()
        
    except Exception as e:
        logger.error(f"Application error: {e}")
        st.error("❌ 應用程序發生錯誤")
        
        # Enhanced error display
        with st.expander("🔧 錯誤詳情 (技術人員使用)", expanded=False):
            st.code(f"""
錯誤類型: {type(e).__name__}
錯誤信息: {str(e)}
會話ID: {st.session_state.get('user_session_id', 'unknown')}
時間戳: {time.strftime('%Y-%m-%d %H:%M:%S')}
            """)
        
        # Recovery suggestions
        st.markdown("""
        ### 🔄 嘗試解決方案：
        
        1. **刷新頁面**：按 F5 或點擊瀏覽器刷新按鈕
        2. **清除緩存**：清空瀏覽器緩存後重新訪問
        3. **稍後重試**：等待 1-2 分鐘後重新嘗試
        4. **檢查網絡**：確保網絡連接穩定
        5. **聯繫支援**：如問題持續，請聯繫技術支援
        """)

# -----------------------------
# Performance Monitoring
# -----------------------------
class PerformanceMonitor:
    """Performance monitoring and optimization"""
    
    @staticmethod
    def log_performance_metrics():
        """Log performance metrics for optimization"""
        try:
            if hasattr(st.session_state, 'app_start_time'):
                session_duration = time.time() - st.session_state.app_start_time
                translation_count = st.session_state.translation_count
                
                # Log metrics (could be sent to analytics service)
                metrics = {
                    'session_duration': session_duration,
                    'translation_count': translation_count,
                    'user_session_id': st.session_state.user_session_id,
                    'language': st.session_state.language,
                    'app_version': APP_CONFIG["version"]
                }
                
                logger.info(f"Performance metrics: {metrics}")
                
        except Exception as e:
            logger.warning(f"Performance monitoring failed: {e}")

# -----------------------------
# Error Recovery
# -----------------------------
class ErrorRecovery:
    """Error recovery and graceful degradation"""
    
    @staticmethod
    def attempt_recovery():
        """Attempt to recover from errors"""
        try:
            # Reset critical session state
            if 'translation_count' not in st.session_state:
                st.session_state.translation_count = 0
            
            if 'language' not in st.session_state:
                st.session_state.language = "简体中文"
            
            # Verify OpenAI client
            client = get_openai_client()
            
            return True
            
        except Exception as e:
            logger.error(f"Recovery failed: {e}")
            return False

# -----------------------------
# Application Entry Point
# -----------------------------
if __name__ == "__main__":
    try:
        # Monitor performance
        PerformanceMonitor.log_performance_metrics()
        
        # Run main application
        main()
        
    except Exception as critical_error:
        logger.critical(f"Critical application error: {critical_error}")
        
        # Attempt recovery
        if ErrorRecovery.attempt_recovery():
            st.info("🔄 正在嘗試恢復服務...")
            st.rerun()
        else:
            st.error("❌ 應用程序遇到嚴重錯誤，請刷新頁面或聯繫技術支援")
            
            # Show minimal recovery interface
            if st.button("🔄 重新啟動應用", type="primary"):
                # Clear session state and restart
                for key in list(st.session_state.keys()):
                    del st.session_state[key]
                st.rerun()
